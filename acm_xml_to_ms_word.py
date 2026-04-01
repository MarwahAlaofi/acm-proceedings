import argparse
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt


def extract_papers_from_xml(xml_file):
    tree = ET.parse(xml_file)
    root = tree.getroot()

    papers = []

    for paper in root.findall("paper"):
        title = paper.findtext("paper_title", default="")

        authors = []
        for author in paper.find("authors").findall("author"):
            first = author.findtext("first_name", "")
            middle = author.findtext("middle_name", "")
            last = author.findtext("last_name", "")

            name = " ".join([first, middle, last]).strip()
            authors.append(name)

        papers.append({
            "title": title,
            "authors": authors
        })

    return papers


def create_word_doc(papers, output_file="papers_list.docx"):
    doc = Document()

    for paper in papers:
        p = doc.add_paragraph()

        # Title with bullet
        run = p.add_run(f"● {paper['title']}")
        run.font.size = Pt(12)

        # Line break
        run.add_break()

        # Authors on next line
        authors_line = ", ".join(paper["authors"])
        run = p.add_run(authors_line)
        run.font.size = Pt(11)

        # spacing after each paper
        p.paragraph_format.space_after = Pt(10)

    doc.save(output_file)
    print(f"Word file generated: {output_file}")


# ----------------------------
# Run
# ----------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Convert ACM XML to Word document"
    )

    parser.add_argument(
        "--input_xml",
        required=True,
        help="Path to input XML file"
    )

    parser.add_argument(
        "--output_docx",
        default="papers_list.docx",
        help="Output Word file name (default: papers_list.docx)"
    )

    args = parser.parse_args()

    papers = extract_papers_from_xml(args.input_xml)
    create_word_doc(papers, args.output_docx)
