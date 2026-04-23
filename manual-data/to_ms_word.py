import argparse
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt

# ----------------------------
# track mapping
# ----------------------------
TRACK_INFO = {
    "Full Paper": {"code": "fp", "heading": "Full Papers"},
    "Short Paper": {"code": "sp", "heading": "Short Papers"},
    "Resource Paper": {"code": "rr", "heading": "Resource Papers"},
    "Reproducibility Paper": {"code": "rp", "heading": "Reproducibility Papers"},
    "Demo Short Paper": {"code": "de", "heading": "Demo Short Papers"},
    "Perspective Paper": {"code": "pr", "heading": "Perspective Papers"},
    "Industry Paper": {"code": "ip", "heading": "Industry Papers"},
    "Tutorial Paper": {"code": "tt", "heading": "Tutorials"},
    "Low Resource Environment": {"code": "lre", "heading": "Low Resource Environment"},
    "Doctoral Abstract": {"code": "dc", "heading": "Doctoral Consortium"},
    "Workshop Summary": {"code": "wk", "heading": "Workshops"},
}

def get_track_info(paper_type):
    return TRACK_INFO.get(paper_type, {"code": "unk", "heading": paper_type})


# ----------------------------
# Extract from XML
# ----------------------------
def extract_papers_from_xml(xml_file):
    tree = ET.parse(xml_file)
    root = tree.getroot()

    papers = []

    for paper in root.findall("paper"):
        title = paper.findtext("paper_title", default="")
        paper_type = paper.findtext("paper_type") or "Unknown"

        authors = []
        for author in paper.find("authors").findall("author"):
            first = author.findtext("first_name", "")
            middle = author.findtext("middle_name", "")
            last = author.findtext("last_name", "")
            name = " ".join([first, middle, last]).strip()
            authors.append(name)

        papers.append({
            "title": title,
            "authors": authors,
            "type": paper_type
        })

    return papers


# ----------------------------
# Create Word doc
# ----------------------------
def create_word_doc(papers, output_file="paper_list.docx"):
    doc = Document()

    # Group papers
    grouped = {}
    for paper in papers:
        grouped.setdefault(paper["type"], []).append(paper)

    # Stats
    print("\n=== Paper Type Counts ===")
    for ptype, plist in grouped.items():
        print(f"{ptype}: {len(plist)}")
    print("=========================\n")

    # Write content of paper types, ordered by type count
    for paper_type, papers_in_type in sorted(
        grouped.items(),
        key=lambda x: len(x[1]),
        reverse=True
    ):
        info = get_track_info(paper_type)
        heading_text = info["heading"]

        # Heading (collapsible)
        doc.add_heading(f"{heading_text} ({len(papers_in_type)})", level=1)

        for paper in papers_in_type:
            info = get_track_info(paper["type"])
            code = info["code"]

            p = doc.add_paragraph()

            # Code (normal)
            run_code = p.add_run(f"[{code}] ")
            run_code.font.size = Pt(12)

            # Title (italic)
            run_title = p.add_run(paper["title"])
            run_title.italic = True
            run_title.font.size = Pt(12)

            # Authors (same paragraph, new line)
            run_auth = p.add_run()
            run_auth.add_break()
            run_auth = p.add_run(", ".join(paper["authors"]))
            run_auth.font.size = Pt(11)

            p.paragraph_format.space_after = Pt(10)

    doc.save(output_file)
    print(f"Word file generated: {output_file}")


# ----------------------------
# Run
# ----------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Convert ACM XML(s) to Word document (simple formatting)"
    )

    parser.add_argument(
        "--input_xml",
        nargs="+",
        required=True,
        help="Path(s) to input XML file(s)"
    )

    parser.add_argument(
        "--output_docx",
        default="papers_list.docx",
        help="Output Word file name"
    )

    args = parser.parse_args()

    all_papers = []

    for xml_file in args.input_xml:
        papers = extract_papers_from_xml(xml_file)
        all_papers.extend(papers)

    create_word_doc(all_papers, args.output_docx)